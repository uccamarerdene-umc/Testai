# -*- coding: utf-8 -*-
import streamlit as st
import os
import glob
import time
import unicodedata
from docx import Document as DocxReader
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_core.documents import Document as LCDocument
from langchain_pinecone import PineconeVectorStore
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pinecone import Pinecone, ServerlessSpec

# 1. Тохиргоо болон Нууцлал
st.set_page_config(
    page_title="Central Test AI Assistant", 
    page_icon="🤖", 
    layout="wide",
    initial_sidebar_state="collapsed"
)

# --- CUSTOM CSS (Figma Design Concept) ---
st.markdown("""
    <style>
    .stApp { background-color: #F8FAFC; }
    h1 { color: #1E293B; font-family: 'Inter', sans-serif; font-weight: 700 !important; }
    section[data-testid="stSidebar"] { background-color: #FFFFFF !important; border-right: 1px solid #E2E8F0; }
    .stButton>button { 
        width: 100%; border-radius: 12px; border: none; 
        background-color: #4F46E5; color: white; padding: 10px 24px; transition: all 0.3s ease; 
    }
    .stButton>button:hover { background-color: #4338CA; transform: translateY(-2px); }
    .stTextInput>div>div>input { border-radius: 12px; border: 1px solid #E2E8F0; padding: 12px; }
    .answer-box { 
        background-color: white; padding: 25px; border-radius: 16px; 
        border: 1px solid #E2E8F0; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1); margin-top: 20px; 
        color: #1E293B; line-height: 1.6;
    }
    </style>
    """, unsafe_allow_html=True)

# API Keys (Cleaned)
def get_key(name):
    val = st.secrets.get(name)
    if val:
        # Тусгай тэмдэгтүүдийг цэвэрлэх
        return str(val).replace('—', '-').strip()
    return None

google_api_key = get_key("GOOGLE_API_KEY")
pinecone_api_key = get_key("PINECONE_API_KEY")
index_name = "testai"

if not google_api_key or not pinecone_api_key:
    st.error("API keys missing! Streamlit Secrets-ээ шалгана уу.")
    st.stop()

@st.cache_resource
def load_models():
    # task_type нь хайлтын чанарыг сайжруулна
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001",
        google_api_key=google_api_key,
        task_type="retrieval_document"
    )
    return embeddings

embeddings = load_models()

# --- DOCUMENT LOADER (Supports Tables) ---
def clean_text(text):
    if not text: return ""
    text = unicodedata.normalize("NFKC", text)
    return "".join(c for c in text if unicodedata.category(c)[0] != 'C' or c in '\n\r\t')

def load_docx_with_tables():
    docs = []
    if not os.path.exists("data"):
        os.makedirs("data")
        return docs
    
    for file in glob.glob("data/*.docx"):
        try:
            doc_obj = DocxReader(file)
            # 1. Текст унших
            paras = [p.text for p in doc_obj.paragraphs if p.text.strip()]
            if paras:
                docs.append(LCDocument(page_content=clean_text("\n".join(paras)), metadata={"source": file}))
            
            # 2. Хүснэгт унших (Markdown хэлбэрээр)
            for table in doc_obj.tables:
                table_data = []
                for row in table.rows:
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    table_data.append(f"| {' | '.join(cells)} |")
                if table_data:
                    docs.append(LCDocument(
                        page_content=f"Хүснэгтийн өгөгдөл ({file}):\n" + "\n".join(table_data),
                        metadata={"source": file}
                    ))
        except Exception as e:
            st.warning(f"Error loading {file}: {e}")
    return docs

# --- SIDEBAR ADMIN (Sync) ---
with st.sidebar:
    st.header("⚙️ Удирдлага")
    admin_pwd = st.text_input("Нууц үг", type="password")
    if admin_pwd == "admin123": # Нууц үг
        if st.button("🔄 Мэдээлэл шинэчлэх (Sync)"):
            all_docs = load_docx_with_tables()
            if not all_docs:
                st.error("Файл олдсонгүй! 'data' хавтсанд .docx файлуудаа хийнэ үү.")
            else:
                with st.spinner("Pinecone-д хадгалж байна. Түр хүлээнэ үү..."):
                    try:
                        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
                        texts = splitter.split_documents(all_docs)
                        
                        pc = Pinecone(api_key=pinecone_api_key)
                        
                        # Хуучин индексийг устгаж шинээр үүсгэх (Dimension 768)
                        if index_name in [i["name"] for i in pc.list_indexes()]:
                            pc.delete_index(index_name)
                            time.sleep(5)
                        
                        pc.create_index(
                            name=index_name, dimension=768, metric="cosine",
                            spec=ServerlessSpec(cloud="aws", region="us-east-1")
                        )
                        
                        # Индекс бэлэн болтол хүлээх
                        while not pc.describe_index(index_name).status['ready']:
                            time.sleep(1)
                        
                        # Batch upload (Google Rate limit-ээс хамгаалах)
                        batch_size = 2 
                        vectorstore = None
                        for i in range(0, len(texts), batch_size):
                            batch = texts[i:i + batch_size]
                            if i == 0:
                                vectorstore = PineconeVectorStore.from_documents(
                                    batch, embeddings, index_name=index_name, pinecone_api_key=pinecone_api_key
                                )
                            else:
                                vectorstore.add_documents(batch)
                            time.sleep(2) # 2 секундын амралт
                            
                        st.success(f"Амжилттай! Нийт {len(texts)} хэсэг хадгалагдлаа.")
                    except Exception as e:
                        st.error(f"Sync алдаа: {e}")

# --- CHAT INTERFACE ---
st.title("🤖 Central Test AI туслах")
st.markdown("<p style='color: #64748B;'>Байгууллагын хүний нөөц, сэтгэл зүйн үнэлгээний ухаалаг туслах</p>", unsafe_allow_html=True)
st.markdown("---")

query = st.text_input("Асуултаа бичнэ үү:", placeholder="Мэдээллийн сангаас хайх...")

if query:
    with st.spinner("Хариулт бэлдэж байна..."):
        try:
            vectorstore = PineconeVectorStore(index_name=index_name, embedding=embeddings, pinecone_api_key=pinecone_api_key)
            search_results = vectorstore.similarity_search(query, k=5)
            
            if search_results:
                context = "\n\n".join([doc.page_content for doc in search_results])
                
                # Gemini 2.5 Flash (Тогтвортой хувилбар)
                llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=google_api_key, temperature=0.1)
                
                prompt = f"""
                Чи бол Central Test компанийн академик зөвлөх, AI туслах юм. 
                Доорх мэдээлэлд тулгуурлан асуултанд ЗӨВХӨН МОНГОЛ хэлээр хариул.
                Эх сурвалж англи хэл дээр байсан ч чи монгол руу маш сайн хөрвүүлж хариулах ёстой.
                Хэрэв мэдээлэлд хүснэгт байгаа бол түүнийг маш цэгцтэй тайлбарла.
                
                Мэдээлэл:
                {context}
                
                Асуулт: {query}
                """
                
                response = llm.invoke(prompt)
                
                st.markdown(f'### 🤖 AI Хариулт:')
                st.markdown(f'<div class="answer-box">{response.content}</div>', unsafe_allow_html=True)
                
                with st.expander("Эх сурвалж харах"):
                    for doc in search_results:
                        st.caption(f"Файл: {doc.metadata.get('source')}")
                        st.text(doc.page_content[:400] + "...")
            else:
                st.warning("⚠️ Холбогдох мэдээлэл олдсонгүй. Мэдээллийн сангаа 'Sync' хийсэн эсэхээ шалгана уу.")
                
        except Exception as e:
            st.error(f"Хайлт хийхэд алдаа гарлаа: {e}")
